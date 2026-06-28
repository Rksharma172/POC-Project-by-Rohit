from docx import Document


def table_to_markdown(table) -> str:
    rows = []

    for row in table.rows:
        cells = [
            cell.text.strip().replace("\n", " ")
            for cell in row.cells
        ]

        if any(cells):
            rows.append(cells)

    if not rows:
        return ""

    width = max(len(row) for row in rows)
    rows = [
        row + [""] * (width - len(row))
        for row in rows
    ]

    header = rows[0]
    separator = ["---"] * width
    body = rows[1:] if len(rows) > 1 else []
    markdown_rows = [header, separator] + body

    return "\n".join(
        "| " + " | ".join(row) + " |"
        for row in markdown_rows
    )


def parse_docx(file_path):
    try:
        doc = Document(file_path)
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        for index, table in enumerate(doc.tables, start=1):
            table_text = table_to_markdown(table)

            if table_text:
                parts.append(
                    f"[TABLE {index}]\n{table_text}\n[/TABLE]"
                )

        text = "\n\n".join(parts)

        if text.strip():
            print("  DOCX parsed successfully")
            return text
        else:
            print("  DOCX appears empty")
            return ""

    except Exception as e:
        print(f"  DOCX parsing failed: {e}")
        return ""
