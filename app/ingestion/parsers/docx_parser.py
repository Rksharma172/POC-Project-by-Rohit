from __future__ import annotations

import asyncio
from pathlib import Path
from typing import Any

from app.core.exceptions import DocumentParseError
from app.core.logging import get_logger

from .base import BaseParser, ParsedDocument

logger = get_logger(__name__)


class DOCXParser(BaseParser):
    """DOCX document parser using python-docx."""

    def supported_extensions(self) -> list[str]:
        return ["docx", "doc"]

    async def parse(self, file_path: Path) -> ParsedDocument:
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self._parse_sync, file_path)

    def _parse_sync(self, file_path: Path) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError as exc:
            raise DocumentParseError("python-docx not installed. Run: pip install python-docx") from exc

        try:
            doc = Document(str(file_path))
            paragraphs: list[str] = []
            sections: list[dict[str, Any]] = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                paragraphs.append(text)
                style_name = para.style.name.lower() if para.style else ""
                if "heading" in style_name:
                    sections.append({"heading": text, "style": para.style.name})

            full_text = "\n".join(paragraphs)
            return ParsedDocument(
                text=full_text,
                sections=sections,
                metadata={
                    "paragraph_count": len(paragraphs),
                    "file_type": "docx",
                    "file_name": file_path.name,
                },
            )
        except Exception as exc:
            logger.error("docx_parse_failed", file=str(file_path), error=str(exc))
            raise DocumentParseError(f"Failed to parse DOCX '{file_path.name}': {exc}") from exc
